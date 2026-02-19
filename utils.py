import os
import fitz
from docx import Document as DocxDocument
from langchain_text_splitters import RecursiveCharacterTextSplitter
from langchain_core.documents import Document
from logger import logger

def extract_text_from_pdf(file_path: str) -> str:
    text = ""
    with fitz.open(file_path) as doc:
        for page in doc:
            text += page.get_text()
    return text

def extract_text_from_docx(file_path: str) -> str:
    doc = DocxDocument(file_path)
    text = "\n".join([para.text for para in doc.paragraphs])
    return text

def chunk_text(text: str, file_path: str = "unknown", chunk_size: int = 1000, chunk_overlap: int = 100) -> list[Document]:
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
        separators=["\n\n", "\n", " ", ""]
    )
    try:
        chunks = text_splitter.split_text(text)
        docs = [Document(page_content=chunk, metadata={"source": file_path}) for chunk in chunks]
        logger.info(f"Chunked {file_path} into {len(docs)} fragments.")
        return docs
    except Exception as e:
        logger.error(f"Error chunking {file_path}: {e}")
        raise

def process_file(file_path: str) -> list[Document]:
    logger.info(f"Starting processing for file: {file_path}")
    ext = os.path.splitext(file_path)[1].lower()
    if ext == ".pdf":
        text = extract_text_from_pdf(file_path)
    elif ext == ".docx":
        text = extract_text_from_docx(file_path)
    else:
        logger.error(f"Unsupported file type: {ext}")
        raise ValueError(f"Unsupported file type: {ext}")
    
    return chunk_text(text, file_path=file_path)
